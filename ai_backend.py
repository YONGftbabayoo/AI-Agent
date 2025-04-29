import os
import openai
import deepl
import docx
import pdfplumber
import pytesseract
from PIL import Image
from googletrans import Translator as GoogleTranslator
from langdetect import detect
from datetime import datetime
from gtts import gTTS
import json
import requests
from bs4 import BeautifulSoup
import spacy
from collections import Counter
import smtplib
from email.message import EmailMessage

# Load spaCy Chinese model (install via: python -m spacy download zh_core_web_sm)
try:
    nlp_zh = spacy.load("zh_core_web_sm")
except:
    nlp_zh = None
    print("[!] Model spaCy 'zh_core_web_sm' belum terinstal. Jalankan: python -m spacy download zh_core_web_sm")

QUALITY_THRESHOLD = 80  # Skor minimal kualitas terjemahan

EMAIL_ADDRESS = "enricosteven611@gmail.com"
EMAIL_PASSWORD = "yvjjtzwgtxvbkchy"

# Konfigurasi API
OPENAI_API_KEY = "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"
DEEPL_API_KEY = "6677f0e1-64a0-47e7-96be-8bda06b84202:fx"
openai.api_key = OPENAI_API_KEY
translator_deepl = deepl.Translator(DEEPL_API_KEY)
translator_google = GoogleTranslator()

custom_dictionary = {
    "社会主义": "sosialisme",
    "文化大革命": "Revolusi Kebudayaan"
}

def detect_language(text):
    try:
        return detect(text)
    except:
        return "zh"

def translate_contextual(text, target_lang="id"):
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a professional Mandarin to Indonesian translator."},
                {"role": "user", "content": f"Translate this Mandarin text into natural Indonesian: {text}"}
            ]
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print("[!] GPT translation failed:", e)
        return None

def translate_literal(text, target_lang="id"):
    try:
        return translator_google.translate(text, dest=target_lang).text
    except Exception as e:
        print("[!] Google literal translation failed:", e)
        return None

def translate_deepl(text, target_lang="ID"):
    try:
        result = translator_deepl.translate_text(text, target_lang=target_lang)
        return result.text
    except Exception as e:
        print("[!] DeepL translation failed:", e)
        return None

def correct_grammar(text):
    if not text:
        return ""
    prompt = f"Anda adalah seorang ahli bahasa Indonesia yang sangat teliti. Tugas Anda adalah memperbaiki grammar dan ejaan dalam teks berikut agar sesuai dengan kaidah Bahasa Indonesia yang benar dan formal. Pastikan untuk tidak mengubah makna atau konteks dari teks ini: {text}"
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print("[!] Grammar correction failed:", e)
        return text

def interactive_correction(text):
    print("\n🗣️ Ingin menggunakan chatbot koreksi? (y/n): ", end="")
    if input().lower() != "y":
        return text

    print("Masukkan instruksi koreksi dalam Bahasa Indonesia (ENTER dua kali untuk selesai):")
    instructions = []
    while True:
        line = input()
        if not line.strip():
            break
        instructions.append(line)
    prompt = f"Kamu adalah editor profesional. Berikut teks yang ingin saya koreksi:\n\n{text}\n\nInstruksi koreksi:\n" + "\n".join(instructions)

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        revised = response.choices[0].message.content.strip()
        print("\n✅ Hasil koreksi dari chatbot:\n", revised[:300], "...\n")
        return revised
    except Exception as e:
        print("[!] Gagal menggunakan chatbot koreksi:", e)
        return text

def summarize_text(text):
    if not text:
        return ""
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": f"Ringkas teks berikut dalam 2-3 kalimat Bahasa Indonesia: {text}"}]
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print("[!] Summarization failed:", e)
        return ""

def classify_topic(text):
    prompt = f"Berikan klasifikasi topik utama dari teks berikut dalam satu kata: budaya, politik, ekonomi, sejarah, teknologi, pendidikan, sosial, atau lainnya.\n\n{text}"
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print("[!] Gagal klasifikasi topik:", e)
        return "tidak diketahui"

def score_translation_quality(original, translation):
    if not translation:
        return "Tidak tersedia"
    try:
        prompt = f"Beri skor kualitas terjemahan dari teks berikut dari 0 hingga 100, serta berikan komentar singkat jika ada kekurangan.\n\nTeks asli:\n{original}\n\nHasil terjemahan:\n{translation}"
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print("[!] Scoring failed:", e)
        return "Tidak tersedia"

def extract_entities(text):
    entities = []
    if nlp_zh:
        doc = nlp_zh(text)
    seen = set()
    for ent in doc.ents:
        key = (ent.text.strip(), ent.label_)
        if key not in seen:
            seen.add(key)
            entities.append({"text": ent.text.strip(), "label": ent.label_})
    return entities

def build_glossary(text, translated_text, path):
    glossary = {}
    if not text or not translated_text:
        return
    try:
        original_words = text.strip().split()
        translated_words = translated_text.strip().split()
        for i in range(min(len(original_words), len(translated_words))):
            glossary[original_words[i]] = translated_words[i]
        with open(path, "w", encoding="utf-8") as f:
            for k, v in glossary.items():
                f.write(f"{k} : {v}\n")
    except Exception as e:
        print("[!] Gagal membangun glossary:", e)

def ensure_term_consistency(text, previous_terms=custom_dictionary):
    if not text:
        return ""
    for term, replacement in previous_terms.items():
        text = text.replace(term, replacement)
    return text

def extract_text_from_image(image_path):
    try:
        image = Image.open(image_path)
        result = pytesseract.image_to_string(image, lang='chi_sim')
        print("🖼️ OCR Result:", result[:200])  # tampilkan sebagian
        return result
    except Exception as e:
        print("[!] OCR failed:", e)
        return ""

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        print("[!] PDF extraction failed:", e)
    return text

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_text_from_url(url):
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.content, "html.parser")
        paragraphs = soup.find_all('p')
        text = "\n".join([p.get_text() for p in paragraphs])
        return text
    except Exception as e:
        print("[!] Gagal mengambil teks dari URL:", e)
        return ""

def send_email_with_attachments(to_email, subject, body, attachments=[]):
    msg = EmailMessage()
    msg["From"] = EMAIL_ADDRESS
    msg["To"] = to_email
    msg["Subject"] = subject
    msg.set_content(body)

    for file_path in attachments:
        try:
            with open(file_path, "rb") as f:
                file_data = f.read()
                file_name = os.path.basename(file_path)
            msg.add_attachment(file_data, maintype="application", subtype="octet-stream", filename=file_name)
        except Exception as e:
            print(f"[!] Gagal menambahkan lampiran {file_path}: {e}")

    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
            smtp.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
            smtp.send_message(msg)
            print(f"📬 Email berhasil dikirim ke: {to_email}")
    except Exception as e:
        print(f"[!] Gagal mengirim email: {e}")


def save_to_docx(original, contextual, literal, corrected, summary, quality_score, entities, topic, path):
    doc = docx.Document()
    doc.add_heading('Translation Result', 0)
    doc.add_heading('Original Text:', level=1)
    doc.add_paragraph(original)
    doc.add_heading('Contextual Translation:', level=1)
    doc.add_paragraph(contextual)
    doc.add_heading('Literal Translation:', level=1)
    doc.add_paragraph(literal)
    doc.add_heading('Corrected Translation:', level=1)
    doc.add_paragraph(corrected)
    doc.add_heading('Summary:', level=1)
    doc.add_paragraph(summary)
    doc.add_heading('Translation Quality Score:', level=1)
    doc.add_paragraph(quality_score)
    doc.add_heading('Topik Utama:', level=1)
    doc.add_paragraph(topic)
    doc.add_heading('Named Entities:', level=1)
    for ent in entities:
        doc.add_paragraph(f"- {ent['text']} ({ent['label']})")
    doc.save(path)
    


def process_file(filepath, mode="gpt"):
    output_name = input("Masukkan nama file output (tanpa ekstensi) atau tekan ENTER untuk otomatis: ").strip()
    format_choice = input("Pilih format output (1 = .docx, 2 = .json, 3 = keduanya): ").strip()

    if filepath.endswith(".docx"):
        text = extract_text_from_docx(filepath)
    elif filepath.endswith(".pdf"):
        text = extract_text_from_pdf(filepath)
    elif filepath.lower().endswith(('.png', '.jpg', '.jpeg')):
        text = extract_text_from_image(filepath)
    elif filepath.startswith("http://") or filepath.startswith("https://"):
        text = extract_text_from_url(filepath)
    else:
        print("[!] Format file tidak dikenali.")
        return

    if not text.strip():
        print("[!] Tidak ada teks yang bisa diproses dari sumber.")
        return

    language = detect_language(text)
    print(f"🌐 Bahasa terdeteksi: {language}")

    print("🔄 Menerjemahkan teks...")
    contextual = translate_contextual(text) or translate_deepl(text) or translate_literal(text) or ""
    literal = translate_literal(text) or ""
    contextual = ensure_term_consistency(contextual)

    print("🔧 Memperbaiki grammar...")
    corrected = correct_grammar(contextual)
    corrected = interactive_correction(corrected)


    print("🧠 Membuat ringkasan...")
    summary = summarize_text(contextual)
    
    print("🗂️ Mengklasifikasikan topik...")
    topic = classify_topic(contextual)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = output_name if output_name else os.path.splitext(os.path.basename(filepath))[0] + "_" + timestamp

    os.makedirs("Result", exist_ok=True)

    json_path = f"Result/{base_name}.json"
    docx_path = f"Result/{base_name}.docx"
    glossary_path = f"Result/glossary_{base_name}.txt"

    print("🔊 Menghasilkan audio dari hasil terjemahan...")
    try:
        tts = gTTS(contextual, lang='id')
        audio_path = f"Result/{base_name}.mp3"
        tts.save(audio_path)
        print(f"✅ File audio disimpan di: {audio_path}")
    except Exception as e:
        print("[!] Gagal membuat audio:", e)

    print("📊 Menilai kualitas terjemahan...")
    quality_score = score_translation_quality(text, contextual)
    
    # Cek apakah perlu peringatan kualitas rendah
    try:
        score_number = int(''.join(filter(str.isdigit, quality_score.split('\n')[0])))
        if score_number < QUALITY_THRESHOLD:
            print(f"⚠️  Peringatan: Skor kualitas terjemahan rendah ({score_number}/100). Periksa kembali hasil terjemahan.")
    except:
        print("⚠️  Tidak dapat memproses skor sebagai angka.")

    print("🔎 Mengekstrak entitas nama...")
    entities = extract_entities(text)

    print("📩 Ingin mengirim hasil ke email? (y/n): ", end="")
    if input().lower() == "y":
        to_email = input("Masukkan alamat email tujuan: ").strip()
        attachments = []
        if format_choice in ["1", "3"]: attachments.append(docx_path)
        if format_choice in ["2", "3"]: attachments.append(json_path)
        if os.path.exists(f"Result/{base_name}.mp3"): attachments.append(f"Result/{base_name}.mp3")
        send_email_with_attachments(to_email, "Hasil Terjemahan Dokumen", "Berikut adalah hasil terjemahan dokumen Anda.", attachments)

    result = {
        "original": text,
        "contextual": contextual,
        "literal": literal,
        "corrected": corrected,
        "summary": summary,
        "quality_score": quality_score,
        "named_entities": entities,
        "topic": topic
    }

    if format_choice in ["2", "3"]:
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"✅ File JSON disimpan di: {json_path}")

    if format_choice in ["1", "3"]:
        save_to_docx(text, contextual, literal, corrected, summary, quality_score, entities, topic, docx_path)
        print(f"✅ File DOCX disimpan di: {docx_path}")

    build_glossary(text, contextual, glossary_path)
    print(f"✅ File glossary disimpan di: {glossary_path}")


def process_batch(files, mode="gpt"):
    for file in files:
        print(f"\n📄 Memproses file: {file}")
        process_file(file, mode)

if __name__ == "__main__":
    files_to_process = [
        "D:/AI-Agent/dok/contoh_dokumen_mandarin_budaya.docx",
        "D:/AI-Agent/dok/contoh_dokumen_mandarin.docx",
        "https://www.chinanews.com.cn/cul/2024/05-12/1001234567.shtml"
    ]
    process_batch(files_to_process)
